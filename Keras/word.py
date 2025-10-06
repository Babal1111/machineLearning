from docx import Document

# Create a new Word document
doc = Document()
doc.add_heading("III. Literature Review: Tabular Synthesis and Analytical Summary", level=1)
doc.add_heading("3.1. Table 1: Synthesis of Empirical Studies on Academic Procrastination and Performance", level=2)

# Define table headers
headers = [
    "S. No.", "Title", "Author(s) detail", "Year of publication",
    "Aim/purpose of the study", "Sample detail (mention country/place name also)",
    "Research method", "Findings", "Limitation / Future recommendations"
]

# Add table with 16 rows (1 header + 15 data rows) and 9 columns
table = doc.add_table(rows=16, cols=9)
table.style = "Table Grid"

# Fill header row
for i, header in enumerate(headers):
    table.cell(0, i).text = header

# Data rows
data = [
    ["1", "The Procrastination Equation: A Meta-Analysis of Correlates", "Steel", "2007",
     "To synthesize research on AP correlates and introduce the TMT framework.",
     "Global sample, diverse populations.", "Meta-Analysis",
     "Confirmed general negative link between procrastination and outcome measures; identified self-regulatory failure as a core mechanism.",
     "Need for longitudinal designs to capture dynamic nature of pacing style."],

    ["2", "Active Procrastination and Academic Performance", "Chu & Choi", "2005",
     "To test the differential relationship between active/passive procrastination and academic performance.",
     "University students, South Korea.", "Correlational Study (PASS scale)",
     "Active procrastination was not negatively correlated with GPA; passive procrastination showed a negative link.",
     "Need to explore mechanisms behind active delay; limited generalizability due to sample size."],

    ["3", "Academic Procrastination, Goal Accomplishment, and Intervention", "Haghbin et al.", "2017",
     "To experimentally test whether goal-setting strategies (SMART goals, implementation intentions) reduce AP.",
     "Undergraduate students (N=177).", "Randomized Quasi-Experiment",
     "Interventions did not significantly reduce AP; baseline AP levels were highly predictive of later goal success.",
     "Suggests simple goal-setting is insufficient; TMT factors must be addressed directly."],

    ["4", "The Moderating Role of Self-Efficacy in Academic Emotion and Procrastination", "Hu et al.", "2024",
     "To analyze the mediating role of self-efficacy between negative emotions (anxiety/depression) and AP.",
     "University students (N=71), China.", "Mediation Analysis",
     "Negative academic emotions increased AP likelihood; Self-efficacy partially mediated this link.",
     "Further research on goal orientation as a potential moderator required."],

    ["5", "Meta-analysis on the Relationship between Academic Performance and Procrastination", "Moradi et al.", "2022",
     "To aggregate existing research and calculate the combined effect size between AP and AP.",
     "96 articles, N=55,477 participants.", "Meta-Analysis",
     "Uncovered a modest negative correlation overall, strongly moderated by type (passive negative, active positive).",
     "Findings emphasize the need for assessment tools that differentiate types of procrastination."],

    ["6", "A Functional Analysis of Procrastination", "Svartdal et al.", "2018",
     "To explore procrastination mechanisms using a behavioral function approach (ABC model).",
     "University students, Norway.", "Functional Analysis (Behavioral Focus)",
     "Identified AP as primarily behavioral delay rooted in failures of the impulsive system.",
     "Model highlights the need for intervention focusing on appetitive and aversive contingencies."],

    ["7", "Temporal Trajectories of Procrastination and Motivation", "Steel & Konig", "2018",
     "To examine motivational failures in a realistic longitudinal design consistent with TMT.",
     "Large correlational dataset (N=7400) and longitudinal study.", "Longitudinal Design",
     "Pacing style reflects a hyperbolic curve; critical self-regulatory skills (attention control, energy regulation) account for 74% of variance.",
     "Confirmed TMT and highlighted the intention-action gap in procrastinators."],

    ["8", "CBT Intervention for Severe Procrastination based on TMT", "Rozental et al.", "2017",
     "To assess the efficacy of guided online CBT targeting TMT components (Expectancy, Value, Impulsiveness).",
     "Participants with severe procrastination (N=150).", "Randomized Controlled Trial (RCT)",
     "Moderate to large effect sizes in reducing procrastination were maintained one year later.",
     "Validated TMT components as targets for clinical intervention."],

    ["9", "Procrastination in Virtual vs. Conventional Students", "Farooq", "2023",
     "To explore the relationship between AP and AP among virtual and conventional university students.",
     "University students (N=200), Pakistan.", "Correlational Study",
     "AP negatively predicted AP; virtual students showed significantly higher AP than conventional students.",
     "Findings suggest environmental (mode of education) factors moderate AP prevalence."],

    ["10", "CBT Intervention Targeting TMT Factors in University Students", "(Anticipated Publication)", "2025",
     "To evaluate a modified CBT intervention focusing on Value, Expectancy, and Impulsivity in a student population.",
     "University students (N=71) with self-reported AP issues.", "Randomized Controlled Trial (RCT)",
     "Significant reduction in AP (d=1.09); improvements noted in Value and Impulsivity, but not Expectancy.",
     "Suggests self-efficacy (Expectancy) is harder to shift than behavioral controls."],

    ["11", "Procrastination Styles and Alcohol Outcomes", "O’Malley et al.", "2017",
     "To identify distinct AP styles and predict their association with alcohol-related problems and GPA.",
     "College undergraduates (N=1106).", "Cluster Analysis (Person-Centered)",
     "Non-procrastination and academic productive procrastination were most adaptive; maladaptive styles predicted poor GPA and higher alcohol risk.",
     "AP is a useful risk indicator for broader maladaptive behaviors."],

    ["12", "Dimensions of University Procrastination in Latin America", "Pichen-Fernández & Turpo Chaparro", "2023",
     "To validate models/scales of university procrastination and their dimensions in Latin-American countries.",
     "University students, Peru.", "Systematic Review/Scale Validation",
     "Identified self-efficacy and procrastination as core dimensions in measurement models.",
     "Emphasizes the centrality of self-efficacy in AP measurement across diverse cultural contexts."],

    ["13", "The Role of Attachment in Academic Success", "(Study on Attachment)", "2016",
     "To determine if procrastination moderates the relationship between attachment styles (anxiety/avoidance) and GPA.",
     "College students.", "Correlational Study (Moderation Analysis)",
     "Procrastination moderates the negative relationship between both attachment anxiety and avoidance and cumulative GPA.",
     "Highlights the interplay of personality traits, emotional regulation, and performance outcomes."],

    ["14", "Factors Causing Academic Procrastination", "Jones & Blankenship", "2020",
     "To analyze the factors most commonly cited by students as causes of AP.",
     "Students (N=70).", "Qualitative/Survey Research",
     "Laziness, poor time management, and fatigue were identified as main reported causes.",
     "While reported causes are surface level, they align with TMT factors (low value, low energy)."],

    ["15", "Self-Efficacy, Anxiety, and Procrastination", "Sirois", "2004",
     "To analyze psychological factors that mediate the link between personality and academic success.",
     "University Students.", "Correlational/Regression",
     "Proposed that self-efficacy and procrastination serve as mediating factors between personality traits and academic success.",
     "Underscores self-efficacy’s role in buffering effects of anxiety."]
]

# Fill table rows
for i, row_data in enumerate(data, start=1):
    for j, value in enumerate(row_data):
        table.cell(i, j).text = value

# Save the document
output_path = "Literature_Review_Table.docx"
doc.save(output_path)

print(f"✅ Word file successfully created: {output_path}")
